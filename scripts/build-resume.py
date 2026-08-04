"""生成雷仰双页详细版个人简历（A4 Word 文档）。

用法：
    set RESUME_SRC=E:\个人简历
    python scripts/build-resume.py
输出：
    E:\个人简历\雷仰个人简历.docx
"""

import os
import sys
import copy

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

SRC = os.environ.get("RESUME_SRC", "E:\\")
OUT = os.path.join(SRC, "雷仰个人简历.docx")
PORTRAIT = os.path.join(SRC, "mmexport1773913557119.png")

# ---------------- 设计令牌 ----------------
TOKENS = {
    "page_width_mm": 210,
    "page_height_mm": 297,
    "margin_top_mm": 14,
    "margin_bottom_mm": 13,
    "margin_left_mm": 16,
    "margin_right_mm": 16,
    "content_width_mm": 178,
    "font_latin": "Calibri",
    "font_east": "微软雅黑",
    "body_size": 10.5,
    "body_color": "262626",
    "muted_color": "5A6572",
    "accent": "1F4E79",
    "name_size": 26,
    "name_color": "17365D",
    "section_size": 13,
    "line_spacing": 1.14,
    "space_after_pt": 2.5,
}


def set_run(run, text=None, size=None, bold=None, color=None, latin=None, east=None):
    if text is not None:
        run.text = text
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = latin or TOKENS["font_latin"]
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), latin or TOKENS["font_latin"])
    rfonts.set(qn("w:hAnsi"), latin or TOKENS["font_latin"])
    rfonts.set(qn("w:eastAsia"), east or TOKENS["font_east"])
    return run


def para(doc, before=0, after=None, align=None, line=None, keep_next=False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(TOKENS["space_after_pt"] if after is None else after)
    if line is not None:
        pf.line_spacing = line
    else:
        pf.line_spacing = TOKENS["line_spacing"]
    if align is not None:
        pf.alignment = align
    if keep_next:
        pf.keep_with_next = True
    return p


def add_bottom_border(p, color=TOKENS["accent"], sz=8, space=4):
    ppr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    ppr.append(pbdr)


def section_heading(doc, text):
    p = para(doc, before=8, after=4, keep_next=True)
    set_run(p.add_run(), text, size=TOKENS["section_size"], bold=True, color=TOKENS["accent"])
    add_bottom_border(p)
    return p


def title_line(doc, left, right, size=10.5):
    p = para(doc, before=2, after=2, keep_next=True)
    set_run(p.add_run(), left, size=size, bold=True, color="262626")
    pf = p.paragraph_format
    pf.tab_stops.add_tab_stop(Cm(TOKENS["content_width_mm"] / 10), WD_ALIGN_PARAGRAPH.RIGHT)
    set_run(p.add_run(), "\t" + right, size=size, bold=False, color=TOKENS["muted_color"])
    return p


def bullet(doc, text, bold_prefix=None, size=None, space_after=2):
    p = doc.add_paragraph(style="List Bullet")
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(space_after)
    pf.line_spacing = TOKENS["line_spacing"]
    pf.left_indent = Cm(0.5)
    pf.first_line_indent = Cm(-0.5)
    if bold_prefix:
        set_run(p.add_run(), bold_prefix, size=size or TOKENS["body_size"], bold=True, color="262626")
    set_run(p.add_run(), text, size=size or TOKENS["body_size"], bold=False, color="333333")
    return p


def set_cell_margins(table, top=80, bottom=80, start=60, end=60):
    tbl = table._tbl
    tblPr = tbl.tblPr
    mar = OxmlElement("w:tblCellMar")
    for tag, val in (("top", top), ("left", start), ("bottom", bottom), ("right", end)):
        el = OxmlElement("w:" + tag)
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tblPr.append(mar)


def no_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement("w:" + edge)
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        borders.append(el)
    tblPr.append(borders)


def set_table_widths(table, widths_cm):
    table.autofit = False
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    total_dxa = int(sum(widths_cm) * 567)
    tblW.set(qn("w:w"), str(total_dxa))
    tblW.set(qn("w:type"), "dxa")
    grid = tbl.find(qn("w:tblGrid"))
    for gc, w in zip(grid.findall(qn("w:gridCol")), widths_cm):
        gc.set(qn("w:w"), str(int(w * 567)))
    for row in table.rows:
        for cell, w in zip(row.cells, widths_cm):
            cell.width = Cm(w)


def add_photo_with_border(cell, path, width_cm):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(path, width=Cm(width_cm))
    ppr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement("w:" + edge)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "6")
        el.set(qn("w:color"), "9AA7B4")
        pbdr.append(el)
    ppr.append(pbdr)


def add_field(p, instr):
    run = p.add_run()
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string("8A97A5")
    run.font.name = TOKENS["font_latin"]
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), TOKENS["font_east"])
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr_el = OxmlElement("w:instrText")
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = instr
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1)
    run._r.append(instr_el)
    run._r.append(fld2)


def build():
    doc = Document()

    # 页面设置
    sec = doc.sections[0]
    sec.page_width = Mm(TOKENS["page_width_mm"])
    sec.page_height = Mm(TOKENS["page_height_mm"])
    sec.top_margin = Mm(TOKENS["margin_top_mm"])
    sec.bottom_margin = Mm(TOKENS["margin_bottom_mm"])
    sec.left_margin = Mm(TOKENS["margin_left_mm"])
    sec.right_margin = Mm(TOKENS["margin_right_mm"])
    sec.header_distance = Mm(8)
    sec.footer_distance = Mm(8)

    # Normal 样式
    normal = doc.styles["Normal"]
    normal.font.name = TOKENS["font_latin"]
    normal.font.size = Pt(TOKENS["body_size"])
    normal.font.color.rgb = RGBColor.from_string(TOKENS["body_color"])
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), TOKENS["font_east"])
    npf = normal.paragraph_format
    npf.space_before = Pt(0)
    npf.space_after = Pt(TOKENS["space_after_pt"])
    npf.line_spacing = TOKENS["line_spacing"]

    # List Bullet 样式
    lb = doc.styles["List Bullet"]
    lb.font.name = TOKENS["font_latin"]
    lb.font.size = Pt(TOKENS["body_size"])
    lb.font.color.rgb = RGBColor.from_string("333333")
    lb._element.rPr.rFonts.set(qn("w:eastAsia"), TOKENS["font_east"])
    lbf = lb.paragraph_format
    lbf.space_before = Pt(0)
    lbf.space_after = Pt(2)
    lbf.line_spacing = TOKENS["line_spacing"]
    lbf.left_indent = Cm(0.5)
    lbf.first_line_indent = Cm(-0.5)

    # ================= 头部 =================
    table = doc.add_table(rows=1, cols=2)
    no_table_borders(table)
    set_cell_margins(table)
    set_table_widths(table, [14.2, 3.6])
    left_cell, right_cell = table.rows[0].cells
    right_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p = left_cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    set_run(p.add_run(), "雷仰", size=TOKENS["name_size"], bold=True, color=TOKENS["name_color"])

    p = left_cell.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    set_run(p.add_run(), "求职意向：硬件研发 / 嵌入式方向实习生", size=12, bold=True, color=TOKENS["accent"])

    p = left_cell.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    set_run(p.add_run(), "电话：16609310328　　", size=10)
    set_run(p.add_run(), "邮箱：lylg0328@163.com", size=10)

    p = left_cell.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    set_run(p.add_run(), "常驻地：浙江义乌　　", size=10)
    set_run(p.add_run(), "出生年月：2005.03", size=10)

    p = left_cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run(), "院校：西北民族大学 · 电气工程学院", size=10)

    add_photo_with_border(right_cell, PORTRAIT, 3.2)

    # ================= 教育背景 =================
    section_heading(doc, "教育背景")
    title_line(doc, "西北民族大学　|　电子信息工程（本科）", "2023.09 – 2027.06")
    p = para(doc, before=1, after=2)
    set_run(p.add_run(), "平均学分绩点：3.27 / 5.0　　", size=TOKENS["body_size"], bold=True, color="262626")
    set_run(p.add_run(), "学号：P231610397", size=TOKENS["body_size"], color=TOKENS["muted_color"])
    p = para(doc, after=2)
    set_run(p.add_run(), "主修课程：", size=TOKENS["body_size"], bold=True)
    set_run(
        p.add_run(),
        "C 语言、电路分析、模拟电子技术、数字电子技术、单片机原理及应用、高频电子线路、信号与系统、数字信号处理、微处理器原理及应用、嵌入式系统、EDA 技术及应用、计算机网络、通信原理",
        size=TOKENS["body_size"],
        color="333333",
    )

    # ================= 实习经历 =================
    section_heading(doc, "实习经历")
    title_line(doc, "北京格物世纪科技有限公司　|　硬件研发部 · 硬件实习生", "2026.07（四周）")
    bullet(doc, "参与硬件研发部日常工作，协助完成嵌入式硬件电路设计、元器件选型与样机调试；根据项目需求绘制原理图并进行 PCB 布局布线，配合工程师完成信号测试与问题定位。")
    bullet(doc, "熟悉企业硬件研发流程与文档规范，掌握从需求评审、方案设计到打样验证的完整工作方法。")

    # ================= 项目经历 =================
    section_heading(doc, "项目经历")

    title_line(doc, "UWB 动态定位水上救援机器人　|　项目负责人", "2024.10 – 2025.02")
    p = para(doc, after=2)
    set_run(p.add_run(), "项目简介：", size=TOKENS["body_size"], bold=True)
    set_run(
        p.add_run(),
        "集成 UWB 厘米级定位与 K230 边缘视觉感知的智能救援无人船，以 STM32 为核心，构建“精准定位—智能识别—自主决策—实时交互”的完整救援流程，实现对落水目标的快速发现、厘米级定位与自主抵近。",
        size=TOKENS["body_size"],
        color="333333",
    )
    bullet(doc, "以 STM32 为核心完成硬件底板 PCB 设计，主导方案选型、原理图设计到 PCB 布局全流程；通过数模/射频区域分割，解决高速数字电路、模拟采集电路与 UWB 射频电路之间的噪声干扰与信号完整性问题。")
    bullet(doc, "负责 UWB 定位模块硬件集成与接口设计，确保射频性能最优；设计多通道 ADC 采集电路，构建高实时性、低功耗的传感器数据采集系统。")
    bullet(doc, "完成姿态、水位等多模态传感器驱动电路、信号调理与硬件调试，应用低通滤波优化数据精度；负责外围功能单元的可调节电源模块设计与接口调试。")
    p = para(doc, after=4)
    set_run(p.add_run(), "项目成果：", size=TOKENS["body_size"], bold=True)
    set_run(
        p.add_run(),
        "国家级大学生创新创业训练计划项目；发表论文《UWB辅助水上救援机器人智能决策方法》（通讯作者）；静态定位水平 RMSE 8.7cm，动态避障成功率 96.7%。",
        size=TOKENS["body_size"],
        color="333333",
    )

    title_line(doc, "智能生活垃圾分类一体机　|　项目负责人", "2025.02 – 2025.07")
    p = para(doc, after=2)
    set_run(p.add_run(), "项目简介：", size=TOKENS["body_size"], bold=True)
    set_run(
        p.add_run(),
        "以 STM32F4 为主控的智能垃圾分类处理系统，采用 K230 识别模块结合 YOLOv11 图像算法对四类垃圾精准识别分拣，分拣准确率达 95% 以上，处理效率每分钟 60 次以上。",
        size=TOKENS["body_size"],
        color="333333",
    )
    bullet(doc, "主导控制主板设计，基于系统功耗、时序及接口需求规划电源与时钟架构，确保核心芯片、传感器获得稳定的运算与控制基础。")
    bullet(doc, "完成多层 PCB 设计，设计电源分配网络与电机驱动等大电流电路，保障系统供电可靠；完成硬件系统级调试与验证。")
    bullet(doc, "设计 ADC 采集与低通滤波处理保障数据精度，完成总线架构设计，为传感器等外设完成 I²C、SPI、UART 等接口布局，确保阻抗控制与等长布线，保障信号完整性与稳定通信。")
    p = para(doc, after=4)
    set_run(p.add_run(), "项目成果：", size=TOKENS["body_size"], bold=True)
    set_run(
        p.add_run(),
        "全国大学生嵌入式芯片与系统设计竞赛全国总决赛三等奖；全国集成电路创新创业大赛西北赛区二等奖。",
        size=TOKENS["body_size"],
        color="333333",
    )

    # ================= 荣誉奖项（第 2 页） =================
    h = section_heading(doc, "荣誉奖项")
    awards = [
        ("2025", "国家级大学生创新创业训练计划项目", "国家级"),
        ("2025", "全国大学生嵌入式芯片与系统设计竞赛", "全国总决赛三等奖"),
        ("2025", "中国机器人及人工智能大赛", "全国总决赛三等奖"),
        ("2025", "全国大学生电子设计大赛", "省级一等奖"),
        ("2025", "全国集成电路创新创业大赛", "西北赛区二等奖"),
        ("2025", "中国大学生计算机设计大赛", "西北赛区三等奖"),
        ("2024", "全国大学生嵌入式芯片与系统设计竞赛", "西北赛区二等奖"),
    ]
    for year, name, level in awards:
        p = para(doc, before=1, after=2)
        set_run(p.add_run(), year + "　", size=TOKENS["body_size"], bold=True, color=TOKENS["accent"])
        set_run(p.add_run(), name, size=TOKENS["body_size"], color="333333")
        set_run(p.add_run(), "　（" + level + "）", size=TOKENS["body_size"], color=TOKENS["muted_color"])

    # ================= 专业技能 =================
    section_heading(doc, "专业技能")
    bullet(
        doc,
        "熟练使用 Altium Designer 21、立创 EDA 等工具，具备高速数字电路、模拟电路及射频电路的 PCB 设计经验；理解信号完整性、电源完整性及电磁兼容性设计原则，能独立完成布局布线、阻抗控制及 Gerber 输出全流程。",
        bold_prefix="PCB 与原理图设计：",
    )
    bullet(
        doc,
        "熟练使用示波器、逻辑分析仪、频谱分析仪等工具进行信号测量、故障定位与硬件调试，具备基本的软件调试能力。",
        bold_prefix="硬件调试：",
    )
    bullet(
        doc,
        "熟练使用 C、Python 编写嵌入式程序，熟悉 KEIL、VS Code、PyCharm 等开发环境；熟悉 STM32、ESP32 等 MCU 应用开发，掌握 EXTI、UART、ADC、DMA、I²C、SPI 等外设。",
        bold_prefix="嵌入式开发：",
    )
    bullet(
        doc,
        "两次担任项目负责人，负责方案选型、任务拆分、进度推进与文档输出，具备跨模块协作与团队沟通能力。",
        bold_prefix="项目管理：",
    )

    # ================= 自我评价 =================
    section_heading(doc, "自我评价")
    p = para(doc, after=2)
    set_run(
        p.add_run(),
        "热爱硬件设计，习惯从系统角度把想法做成能稳定运行的板子；动手能力强、学习新工具快，注重电路细节与工程规范，愿意扎根硬件研发一线持续打磨能力，期待在嵌入式硬件方向长期发展。",
        size=TOKENS["body_size"],
        color="333333",
    )

    # ================= 页脚 =================
    footer = sec.footer
    fp = footer.paragraphs[0]
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    fp.paragraph_format.tab_stops.add_tab_stop(Cm(17.8), WD_ALIGN_PARAGRAPH.RIGHT)
    set_run(fp.add_run(), "雷仰 · 个人简历", size=8.5, color="8A97A5")
    set_run(fp.add_run(), "\t第 ", size=8.5, color="8A97A5")
    add_field(fp, " PAGE ")
    set_run(fp.add_run(), " 页 / 共 ", size=8.5, color="8A97A5")
    add_field(fp, " NUMPAGES ")
    set_run(fp.add_run(), " 页", size=8.5, color="8A97A5")

    doc.save(OUT)
    print("SAVED:", OUT)


if __name__ == "__main__":
    build()
